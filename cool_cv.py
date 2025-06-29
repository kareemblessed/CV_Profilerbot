import logging
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, CallbackQueryHandler, MessageHandler, filters, ContextTypes, ConversationHandler
import google.generativeai as genai

# Import everything from flow.py
from flow import start, help_command, button_callback, BOT_TOKEN

# Enable logging
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Google Gemini API key - FREE! Get from ai.google.dev
GEMINI_API_KEY = ""  # PUT YOUR ACTUAL API KEY HERE

# Initialize Gemini
genai.configure(api_key=GEMINI_API_KEY)

# Try multiple models to find one that works
def initialize_gemini_model():
    """Try different Gemini models until one works"""
    model_names = [
        'gemini-2.0-flash-exp', 'gemini-1.5-flash', 'gemini-2.0-flash',
        'models/gemini-1.5-flash', 'models/gemini-1.5-pro', 'models/gemini-pro'
    ]
    
    for model_name in model_names:
        try:
            print(f"🔄 Trying model: {model_name}")
            test_model = genai.GenerativeModel(model_name)
            test_response = test_model.generate_content("Hello")
            
            if test_response and test_response.text:
                print(f"✅ Success! Using model: {model_name}")
                return test_model
        except Exception as e:
            print(f"❌ Model {model_name} failed: {str(e)[:100]}")
            continue
    
    raise Exception("❌ No Gemini models are working. Please check your API key at https://ai.google.dev")

# Initialize the working model
print("🔍 Finding a working Gemini model...")
model = initialize_gemini_model()

# Enhanced conversation states
NAME, PHONE, EMAIL, LOCATION, JOB_TITLE, SKILLS, EXPERIENCE_COUNT, EXPERIENCE_ENTRY, EDUCATION_COUNT, EDUCATION_ENTRY, ADDITIONAL = range(11)

def enhance_with_gemini(prompt_type, content, job_title=""):
    """Enhanced Gemini prompts to improve CV content"""
    try:
        logger.info(f"🤖 Calling Gemini API for {prompt_type}...")
        
        prompts = {
            "skills": f"You are a professional CV writer. Generate 8-10 relevant skills for a {job_title}. Based on: {content}. Include technical and soft skills. Format as comma-separated list. Return only the skills list.",
            
            "experience": f"You are a professional CV writer. Enhance this work experience for a {job_title}: {content}. Format as: **Job Title | Company | Dates** followed by 3-4 bullet points with quantified achievements using action verbs. Make it impactful with KPIs and percentages where possible.",
            
            # FIXED EDUCATION PROMPT - Now matches experience structure
            "education": f"You are a professional CV writer. Enhance this education entry for a {job_title}: {content}. Format as: **Degree | Institution | Year** followed by 3-4 bullet points highlighting relevant achievements, projects, GPA (if high), honors, coursework, or academic accomplishments. Make it impactful with specific details, leadership roles, and quantified achievements where possible.",
            
            "summary": f"Create a compelling 3-sentence professional summary for a {job_title}. Based on: {content}. Highlight key strengths, experience years, and value proposition. Return only the summary.",
            
            "additional": f"Format this additional information as clean bullet points: {content}. Create maximum 5 concise bullet points covering certifications, languages, technical proficiencies, or achievements. Each bullet should be one line only."
        }
        
        if prompt_type not in prompts:
            return content
        
        response = model.generate_content(prompts[prompt_type])
        enhanced_text = response.text.strip()
        logger.info(f"✅ Gemini enhanced {prompt_type}")
        return enhanced_text
        
    except Exception as e:
        logger.error(f"❌ Gemini enhancement failed for {prompt_type}: {e}")
        return content

def parse_bold_text(text):
    """Parse text with **bold** markers and return formatted parts"""
    parts = []
    current_text = ""
    is_bold = False
    i = 0
    
    while i < len(text):
        if i < len(text) - 1 and text[i:i+2] == "**":
            if current_text:
                parts.append((current_text, is_bold))
                current_text = ""
            is_bold = not is_bold
            i += 2
        else:
            current_text += text[i]
            i += 1
    
    if current_text:
        parts.append((current_text, is_bold))
    
    return parts

def add_formatted_paragraph(doc, text):
    """Add a paragraph with proper bold formatting"""
    para = doc.add_paragraph()
    parts = parse_bold_text(text)
    
    for text_part, is_bold in parts:
        run = para.add_run(text_part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        if is_bold:
            run.bold = True
    
    return para

def create_cv_document(user_data):
    """Create a professional CV document with AI enhancement"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)
    
    # NAME (Large, Bold, Centered)
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(user_data.get('name', '').upper())
    name_run.font.name = 'Times New Roman'
    name_run.font.size = Pt(22)
    name_run.bold = True
    
    # JOB TITLE (Italic, Centered)
    if user_data.get('job_title'):
        job_para = doc.add_paragraph()
        job_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        job_run = job_para.add_run(user_data.get('job_title'))
        job_run.font.name = 'Times New Roman'
        job_run.font.size = Pt(14)
        job_run.italic = True
    
    # Line separator
    separator_para = doc.add_paragraph()
    separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    separator_run = separator_para.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    separator_run.font.name = 'Times New Roman'
    separator_run.font.size = Pt(10)
    
    # CONTACT INFO (Centered)
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info = []
    
    if user_data.get('phone'):
        contact_info.append(f"📱 {user_data.get('phone')}")
    if user_data.get('email'):
        contact_info.append(f"✉️ {user_data.get('email')}")
    if user_data.get('location'):
        contact_info.append(f"📍 {user_data.get('location')}")
    
    contact_run = contact_para.add_run(' | '.join(contact_info))
    contact_run.font.name = 'Times New Roman'
    contact_run.font.size = Pt(11)
    doc.add_paragraph()
    
    # PROFESSIONAL SUMMARY
    if user_data.get('skills') and user_data.get('job_title'):
        summary_content = f"Job title: {user_data.get('job_title')}. Skills: {user_data.get('skills')}. Experience: {user_data.get('experiences', [''])[0] if user_data.get('experiences') else ''}."
        enhanced_summary = enhance_with_gemini("summary", summary_content, user_data.get('job_title'))
        
        summary_heading = doc.add_paragraph()
        summary_heading_run = summary_heading.add_run("PROFESSIONAL SUMMARY")
        summary_heading_run.bold = True
        summary_heading_run.font.name = 'Times New Roman'
        summary_heading_run.font.size = Pt(14)
        summary_heading_run.font.color.rgb = RGBColor(31, 78, 121)
        
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(enhanced_summary.replace('**', ''))
        summary_run.font.name = 'Times New Roman'
        summary_run.font.size = Pt(11)
        doc.add_paragraph()
    
    # CORE COMPETENCIES (Two-column layout)
    if user_data.get('skills'):
        enhanced_skills = enhance_with_gemini("skills", user_data.get('skills'), user_data.get('job_title', ''))
        
        skills_heading = doc.add_paragraph()
        skills_heading_run = skills_heading.add_run("CORE COMPETENCIES")
        skills_heading_run.bold = True
        skills_heading_run.font.name = 'Times New Roman'
        skills_heading_run.font.size = Pt(14)
        skills_heading_run.font.color.rgb = RGBColor(31, 78, 121)
        
        skills_list = [skill.strip() for skill in enhanced_skills.split(',')]
        mid_point = (len(skills_list) + 1) // 2
        left_column = skills_list[:mid_point]
        right_column = skills_list[mid_point:]
        
        table = doc.add_table(rows=max(len(left_column), len(right_column)), cols=2)
        table.columns[0].width = Pt(250)
        table.columns[1].width = Pt(250)
        table.style = 'Light List'
        
        for i in range(max(len(left_column), len(right_column))):
            row = table.rows[i]
            
            if i < len(left_column):
                cell = row.cells[0]
                cell.text = f"• {left_column[i]}"
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.font.bold = False  # This makes sure the text is not bold
                        run.font.italic = True  # This makes the text italic
            
            if i < len(right_column):
                cell = row.cells[1]
                cell.text = f"• {right_column[i]}"
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.font.bold = False  # This makes sure the text is not bold
                        run.font.italic = True  # This makes the text italic
        
        doc.add_paragraph()
    
    # PROFESSIONAL EXPERIENCE (Multiple entries with bold formatting)
    if user_data.get('experiences'):
        exp_heading = doc.add_paragraph()
        exp_heading_run = exp_heading.add_run("PROFESSIONAL EXPERIENCE")
        exp_heading_run.bold = True
        exp_heading_run.font.name = 'Times New Roman'
        exp_heading_run.font.size = Pt(14)
        exp_heading_run.font.color.rgb = RGBColor(31, 78, 121)
        
        for experience in user_data.get('experiences', []):
            enhanced_experience = enhance_with_gemini("experience", experience, user_data.get('job_title', ''))
            
            # Split into lines and process
            exp_lines = enhanced_experience.split('\n')
            for line in exp_lines:
                if line.strip():
                    add_formatted_paragraph(doc, line.strip())
            
            doc.add_paragraph()  # Space between experiences
    
    # EDUCATION (Multiple entries with bold formatting) - FIXED TO MATCH EXPERIENCE
    if user_data.get('education_entries'):
        edu_heading = doc.add_paragraph()
        edu_heading_run = edu_heading.add_run("EDUCATION")
        edu_heading_run.bold = True
        edu_heading_run.font.name = 'Times New Roman'
        edu_heading_run.font.size = Pt(14)
        edu_heading_run.font.color.rgb = RGBColor(31, 78, 121)
        
        for education in user_data.get('education_entries', []):
            # FIXED: Now passes job_title like experience does
            enhanced_education = enhance_with_gemini("education", education, user_data.get('job_title', ''))
            
            # FIXED: Now processes exactly like experience (removed the startswith filter)
            edu_lines = enhanced_education.split('\n')
            for line in edu_lines:
                if line.strip():  # Only check if line has content, just like experience
                    add_formatted_paragraph(doc, line.strip())
            
            doc.add_paragraph()  # Space between education entries
    
    # ADDITIONAL QUALIFICATIONS (Clean bullet points)
    if user_data.get('additional') and user_data.get('additional').lower() != 'skip':
        enhanced_additional = enhance_with_gemini("additional", user_data.get('additional'))
        
        additional_heading = doc.add_paragraph()
        additional_heading_run = additional_heading.add_run("ADDITIONAL QUALIFICATIONS")
        additional_heading_run.bold = True
        additional_heading_run.font.name = 'Times New Roman'
        additional_heading_run.font.size = Pt(14)
        additional_heading_run.font.color.rgb = RGBColor(31, 78, 121)
        
        # Process additional info as bullet points (max 5)
        additional_lines = enhanced_additional.split('\n')
        bullet_count = 0
        for line in additional_lines:
            if line.strip() and bullet_count < 5:
                # Remove existing bullets and add our own
                clean_line = line.strip().lstrip('•').lstrip('-').lstrip('*').strip()
                if clean_line:
                    bullet_para = doc.add_paragraph()
                    bullet_run = bullet_para.add_run(f"• {clean_line}")
                    bullet_run.font.name = 'Times New Roman'
                    bullet_run.font.size = Pt(11)
                    bullet_count += 1
    
    return doc

# Button callback for choosing number of entries
async def experience_count_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle experience count selection"""
    query = update.callback_query
    await query.answer()
    
    count = int(query.data.split('_')[1])
    context.user_data['exp_count'] = count
    context.user_data['exp_current'] = 1
    context.user_data['experiences'] = []
    
    await query.edit_message_text(
        f"💼 Experience Entry 1 of {count}\n\n"
        "📝 Please provide your MOST RECENT work experience:\n\n"
        "Include: Job title, Company, Dates, and key achievements\n"
        "Example: Senior Developer at TechCorp, 2022-Present, Led team of 8, increased efficiency by 40%"
    )
    return EXPERIENCE_ENTRY

async def education_count_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle education count selection"""
    query = update.callback_query
    await query.answer()
    
    count = int(query.data.split('_')[1])
    context.user_data['edu_count'] = count
    context.user_data['edu_current'] = 1
    context.user_data['education_entries'] = []
    
    await query.edit_message_text(
        f"🎓 Education Entry 1 of {count}\n\n"
        "📝 Please provide your MOST RECENT education:\n\n"
        "Include: Degree, Institution, Year, and achievements\n"
        "Example: Master of Computer Science, MIT, 2020, GPA 3.8, Dean's List, Thesis on AI"
    )
    return EDUCATION_ENTRY

# EXTENDED BUTTON CALLBACK
async def extended_button_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Extended button callback that handles CV creation"""
    query = update.callback_query
    
    # Handle experience count selection
    if query.data.startswith('exp_'):
        return await experience_count_callback(update, context)
    
    # Handle education count selection
    if query.data.startswith('edu_'):
        return await education_count_callback(update, context)
    
    # First, try to handle with flow.py button_callback
    try:
        result = await button_callback(update, context)
        if result is not None:
            return result
    except:
        pass
    
    # Handle the 'begin_cv' callback
    if query.data == 'begin_cv':
        await query.answer()
        logger.info("🚀 BEGIN_CV_CREATION called!")
        context.user_data.clear()
        
        await query.edit_message_text("👋 Let's start! What's your full name?")
        return NAME

# CV Creation Conversation Handlers
async def collect_name(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect user's name"""
    context.user_data['name'] = update.message.text
    await update.message.reply_text("📱 What's your phone number?")
    return PHONE

async def collect_phone(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect user's phone"""
    context.user_data['phone'] = update.message.text
    await update.message.reply_text("📧 What's your email address?")
    return EMAIL

async def collect_email(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect user's email"""
    context.user_data['email'] = update.message.text
    await update.message.reply_text("📍 What's your location? (City, Country)")
    return LOCATION

async def collect_location(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect user's location"""
    context.user_data['location'] = update.message.text
    await update.message.reply_text("💼 What's your current/target job title?")
    return JOB_TITLE

async def collect_job_title(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect job title"""
    context.user_data['job_title'] = update.message.text
    await update.message.reply_text("🛠️ List your key skills (comma-separated):\n\nExample: Python, JavaScript, Project Management, Communication")
    return SKILLS

async def collect_skills(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect skills and ask for experience count"""
    context.user_data['skills'] = update.message.text
    
    keyboard = [
        [InlineKeyboardButton("1️⃣", callback_data='exp_1')],
        [InlineKeyboardButton("2️⃣", callback_data='exp_2')],
        [InlineKeyboardButton("3️⃣", callback_data='exp_3')]
    ]
    
    await update.message.reply_text(
        "💼 How many work experiences would you like to add?\n\n"
        "Choose between 1-3 (we'll start with most recent):",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    return EXPERIENCE_COUNT

async def collect_experience_entry(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect individual experience entries"""
    context.user_data['experiences'].append(update.message.text)
    context.user_data['exp_current'] += 1
    
    if context.user_data['exp_current'] <= context.user_data['exp_count']:
        # Ask for next experience (going from recent to older)
        ordinal = ["", "first", "second", "third"][context.user_data['exp_current']]
        await update.message.reply_text(
            f"💼 Experience Entry {context.user_data['exp_current']} of {context.user_data['exp_count']}\n\n"
            f"📝 Please provide your {ordinal} work experience:\n"
            "(Going from most recent to oldest)"
        )
        return EXPERIENCE_ENTRY
    
    # Move to education
    keyboard = [
        [InlineKeyboardButton("1️⃣", callback_data='edu_1')],
        [InlineKeyboardButton("2️⃣", callback_data='edu_2')],
        [InlineKeyboardButton("3️⃣", callback_data='edu_3')]
    ]
    
    await update.message.reply_text(
        "🎓 How many education entries would you like to add?\n\n"
        "Choose between 1-3 (we'll start with most recent):",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )
    return EDUCATION_COUNT

async def collect_education_entry(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect individual education entries"""
    context.user_data['education_entries'].append(update.message.text)
    context.user_data['edu_current'] += 1
    
    if context.user_data['edu_current'] <= context.user_data['edu_count']:
        # Ask for next education
        ordinal = ["", "first", "second", "third"][context.user_data['edu_current']]
        await update.message.reply_text(
            f"🎓 Education Entry {context.user_data['edu_current']} of {context.user_data['edu_count']}\n\n"
            f"📝 Please provide your {ordinal} education:\n"
            "(Going from most recent to oldest)"
        )
        return EDUCATION_ENTRY
    
    # Move to additional info
    await update.message.reply_text(
        "➕ Any additional information?\n\n"
        "Include certifications, languages, technical skills, or achievements.\n"
        "Maximum 5 items will be used.\n\n"
        "Type 'skip' if you don't have any additional info."
    )
    return ADDITIONAL

async def collect_additional(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Collect additional info and generate CV"""
    additional = update.message.text
    context.user_data['additional'] = '' if additional.lower() == 'skip' else additional
    
    await update.message.reply_text("⏳ Generating your professional CV with AI enhancements...\n\n🤖 Creating multiple sections with bold formatting...")
    
    try:
        doc = create_cv_document(context.user_data)
        
        filename = f"CV_{context.user_data.get('name', 'user').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(filename)
        
        with open(filename, 'rb') as cv_file:
            await update.message.reply_document(
                cv_file, 
                caption="✅ Your enhanced CV is ready! 🎉\n\n🔥 Features:\n• Multiple experience entries\n• Multiple education entries\n• Bold job titles & degrees\n• KPIs and achievements\n• Clean bullet points\n• Professional formatting"
            )
        
        os.remove(filename)
        context.user_data.clear()
        
        keyboard = [
            [InlineKeyboardButton("🚀 Create Another CV", callback_data='start_cv')],
            [InlineKeyboardButton("🏠 Main Menu", callback_data='main_menu')]
        ]
        
        await update.message.reply_text(
            "🎉 Enhanced CV created successfully!\n\nWhat would you like to do next?",
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        
    except Exception as e:
        logger.error(f"❌ CV generation error: {e}")
        await update.message.reply_text(f"❌ Error: {str(e)}\n\nPlease try again with /start")
    
    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Cancel conversation"""
    keyboard = [[InlineKeyboardButton("🏠 Main Menu", callback_data='main_menu')]]
    await update.message.reply_text("❌ CV creation cancelled.", reply_markup=InlineKeyboardMarkup(keyboard))
    return ConversationHandler.END

def main():
    """Start the enhanced bot"""
    print("🤖 Enhanced ProfileBot is starting...")
    print("✨ Features: Multiple entries, Bold formatting, Clean bullets")
    
    try:
        print("🔍 Testing Gemini API connection...")
        test_response = model.generate_content("Hello")
        if test_response and test_response.text:
            print("✅ Gemini API is working!")
        else:
            print("⚠️ Gemini API might not be configured correctly")
    except Exception as e:
        print(f"❌ Gemini API test failed: {e}")
        return
    
    try:
        application = Application.builder().token(BOT_TOKEN).build()
        
        conv_handler = ConversationHandler(
            entry_points=[CallbackQueryHandler(extended_button_callback, pattern=r'^begin_cv$')],
            states={
                NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, collect_name)],
                PHONE: [MessageHandler(filters.TEXT & ~filters.COMMAND, collect_phone)],
                EMAIL: [MessageHandler(filters.TEXT & ~filters.COMMAND, collect_email)],
                LOCATION: [MessageHandler(filters.TEXT & ~filters.COMMAND, collect_location)],
                JOB_TITLE: [MessageHandler(filters.TEXT & ~filters.COMMAND, collect_job_title)],
                SKILLS: [MessageHandler(filters.TEXT & ~filters.COMMAND, collect_skills)],
                EXPERIENCE_COUNT: [CallbackQueryHandler(extended_button_callback, pattern=r'^exp_[1-3]$')],
                EXPERIENCE_ENTRY: [MessageHandler(filters.TEXT & ~filters.COMMAND, collect_experience_entry)],
                EDUCATION_COUNT: [CallbackQueryHandler(extended_button_callback, pattern=r'^edu_[1-3]$')],
                EDUCATION_ENTRY: [MessageHandler(filters.TEXT & ~filters.COMMAND, collect_education_entry)],
                ADDITIONAL: [MessageHandler(filters.TEXT & ~filters.COMMAND, collect_additional)],
            },
            fallbacks=[CommandHandler('cancel', cancel)],
            per_message=False
        )
        
        application.add_handler(CommandHandler("start", start))
        application.add_handler(CommandHandler("help", help_command))
        application.add_handler(conv_handler)
        application.add_handler(CallbackQueryHandler(extended_button_callback))
        
        print("✅ Enhanced bot initialized successfully!")
        print("🚀 Features: 1-3 entries per section, bold formatting, clean design")
        print("🎓 Education section now works exactly like Experience section!")
        print("-" * 50)
        
        application.run_polling(allowed_updates=Update.ALL_TYPES)
        
    except Exception as e:
        print(f"❌ Error starting bot: {e}")

if __name__ == '__main__':
    main()
